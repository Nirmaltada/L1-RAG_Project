import asyncio
import json
import logging
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from pypdf import PdfReader

from app.models import DocumentRecord
from app.providers.groq import GroqProvider
from app.rag.embeddings import EmbeddingService
from app.rag.prompts import CLASSIFICATION_SYSTEM_PROMPT, classification_user_prompt
from app.rag.retrieval import ChromaStores

logger = logging.getLogger(__name__)
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class IngestionError(ValueError):
    pass


def _safe_list(value: Any, limit: int) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip()[:100] for item in value if str(item).strip()][:limit]


def _metadata_defaults() -> dict[str, Any]:
    return {"category": "general", "document_type": "document", "topics": [], "keywords": []}


def parse_file(path: Path) -> tuple[list[Document], str]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise IngestionError("Supported file types are PDF, DOCX, TXT, and Markdown")
    try:
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            title = str((reader.metadata or {}).get("/Title", "") or "")
            documents = []
            for page_number, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    documents.append(Document(text=text, metadata={"page_number": page_number}))
            return documents, title
        if suffix == ".docx":
            source = DocxDocument(str(path))
            text = "\n".join(paragraph.text for paragraph in source.paragraphs if paragraph.text.strip())
            title = source.core_properties.title or ""
            return ([Document(text=text, metadata={})] if text.strip() else []), title
        text = path.read_text(encoding="utf-8-sig").strip()
        return ([Document(text=text, metadata={})] if text else []), ""
    except IngestionError:
        raise
    except Exception as exc:
        raise IngestionError(f"Could not parse {path.name}: {exc}") from exc


class IngestionService:
    def __init__(
        self,
        stores: ChromaStores,
        embeddings: EmbeddingService,
        groq: GroqProvider,
        classification_model: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> None:
        self.stores = stores
        self.embeddings = embeddings
        self.groq = groq
        self.classification_model = classification_model
        self.splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    async def classify(self, filename: str, text: str) -> dict[str, Any]:
        try:
            result = await self.groq.json_completion(
                self.classification_model,
                CLASSIFICATION_SYSTEM_PROMPT,
                classification_user_prompt(filename, text),
            )
            return {
                "category": str(result.get("category") or "general").strip()[:80],
                "document_type": str(result.get("document_type") or "document").strip()[:80],
                "topics": _safe_list(result.get("topics"), 8),
                "keywords": _safe_list(result.get("keywords"), 12),
            }
        except Exception:
            logger.exception("Document metadata classification failed for %s", filename)
            return _metadata_defaults()

    async def ingest(
        self,
        chat_id: str,
        document_id: str,
        original_filename: str,
        stored_filename: str,
        file_hash: str,
        path: Path,
    ) -> DocumentRecord:
        documents, title = await asyncio.to_thread(parse_file, path)
        if not documents:
            raise IngestionError("The document contains no extractable text")
        representative = "\n\n".join(document.text for document in documents)[:12000]
        semantic = await self.classify(original_filename, representative)
        created_at = datetime.now(UTC)
        inherited = {
            "chat_id": chat_id,
            "document_id": document_id,
            "filename": original_filename,
            "file_type": path.suffix.lower().lstrip("."),
            "document_type": semantic["document_type"],
            "category": semantic["category"],
            "topics": json.dumps(semantic["topics"], ensure_ascii=False),
            "keywords": json.dumps(semantic["keywords"], ensure_ascii=False),
            "created_at": created_at.isoformat(),
            "document_title": title,
        }
        for document in documents:
            page_number = document.metadata.get("page_number")
            document.metadata = {**inherited, "page_number": page_number or 0}

        nodes = await asyncio.to_thread(
            self.splitter.get_nodes_from_documents, documents, show_progress=False
        )
        if not nodes:
            raise IngestionError("The document produced no text chunks")
        for node in nodes:
            node.id_ = str(uuid.uuid4())
            node.metadata["chunk_id"] = node.id_
        embeddings, embedding_model = await self.embeddings.embed(
            [node.get_content() for node in nodes],
            expected_dimension=self.stores.embedding_dimension(),
        )
        logger.info("Document %s embedded with model=%s", original_filename, embedding_model)
        if len(embeddings) != len(nodes):
            raise IngestionError("Embedding provider returned the wrong number of vectors")
        for node, embedding in zip(nodes, embeddings, strict=True):
            node.embedding = embedding
        await asyncio.to_thread(self.stores.add_nodes, nodes)
        return DocumentRecord(
            id=document_id,
            chat_id=chat_id,
            filename=original_filename,
            stored_filename=stored_filename,
            file_hash=file_hash,
            file_type=path.suffix.lower().lstrip("."),
            category=semantic["category"],
            document_type=semantic["document_type"],
            topics=semantic["topics"],
            keywords=semantic["keywords"],
            created_at=created_at,
            chunk_count=len(nodes),
        )
